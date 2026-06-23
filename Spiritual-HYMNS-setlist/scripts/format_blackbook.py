from __future__ import annotations

import re
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"D:\worship-song-library")
SOURCE = ROOT / "docs" / "blackbook" / "BLACKBOOK_FULL_SUBMISSION3.docx"
OUTPUT = ROOT / "docs" / "blackbook" / "BLACKBOOK_FINAL_FORMATTED.docx"
DIAGRAMS = ROOT / "blackbook-diagrams"


MANDATORY_IMAGES = {
    "System Architecture": ("system_architecture.png", "Figure 5.1: System architecture of the Worship Song Library."),
    "ER Diagram": ("er_diagram.png", "Figure 6.1: Entity relationship diagram of the Worship Song Library database."),
    "Chord Mapping": ("chord_mapping.png", "Figure 7.1: Chord-to-text mapping using character index positions."),
    "Rendering Flow": ("rendering_flow.png", "Figure 8.1: Rendering flow from structured data to responsive song display."),
}

OPTIONAL_IMAGES = {
    "Chord Screenshot": ("song_display.png", "Figure 7.2: Song display view showing aligned chords and lyrics on mobile."),
    "Alignment Comparison": ("alignment_comparison.png", "Figure 8.2: Comparison of incorrect static alignment and corrected responsive alignment."),
    "Responsive Mobile": ("mobile_responsive.png", "Figure 12.1: Mobile responsive song view of the final application."),
}


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def set_paragraph_border(paragraph, position: str = "bottom", size: int = 12, color: str = "000000", space: int = 1) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    existing = p_bdr.find(qn(f"w:{position}"))
    if existing is not None:
        p_bdr.remove(existing)

    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)
    p_bdr.append(border)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def apply_page_border(section) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgBorders"))
    if existing is not None:
        sect_pr.remove(existing)

    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    pg_borders.set(qn("w:display"), "allPages")
    pg_borders.set(qn("w:zOrder"), "front")

    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:space"), "20")
        border.set(qn("w:color"), "000000")
        pg_borders.append(border)

    sect_pr.append(pg_borders)


def remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def clear_paragraph(paragraph) -> None:
    paragraph.clear()
    paragraph.text = ""


def set_run_font(run, bold: bool | None = None, size: int | None = None) -> None:
    run.font.name = "Times New Roman"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clean_text(text: str) -> str:
    text = text.replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def is_system_marker(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    marker_patterns = [
        r"^\[(?:PAGE BORDER:.*|CHAPTER\s+\w+.*|DIAGRAM DESCRIPTION.*)\]$",
        r"^\[(?:COLLEGE NAME|UNIVERSITY NAME|ROLL NUMBER|GUIDE NAME|GUIDE DESIGNATION|STUDENT NAME)\]$",
    ]
    return any(re.match(pat, stripped, flags=re.I) for pat in marker_patterns)


def style_document(doc: Document) -> None:
    for section in doc.sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.55)
        section.right_margin = Inches(1.05)
        section.header_distance = Inches(0.45)
        section.footer_distance = Inches(0.45)
        apply_page_border(section)
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        footer_p = footer.paragraphs[0]
        clear_paragraph(footer_p)
        add_page_field(footer_p)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(8)
    pf.space_before = Pt(0)
    pf.widow_control = True

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Times New Roman"
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    h1pf = heading1.paragraph_format
    h1pf.line_spacing = 1.5
    h1pf.space_before = Pt(16)
    h1pf.space_after = Pt(8)
    h1pf.keep_with_next = True

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Times New Roman"
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    h2pf = heading2.paragraph_format
    h2pf.line_spacing = 1.5
    h2pf.space_before = Pt(12)
    h2pf.space_after = Pt(6)
    h2pf.keep_with_next = True


def normalize_paragraphs(doc: Document) -> None:
    for para in list(doc.paragraphs):
        text = para.text or ""
        if is_system_marker(text):
            remove_paragraph(para)
            continue

        if para.style.name == "Normal":
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.widow_control = True

        stripped = text.strip()
        if not stripped:
            continue

        if para.style.name == "Heading 1":
            match = re.match(r"Chapter\s+(\d+):\s*(.+)", stripped, re.I)
            if match:
                para.text = f"Chapter {match.group(1)}: {match.group(2).strip()}"
            else:
                para.text = stripped
        elif para.style.name == "Heading 2":
            para.text = stripped
        else:
            cleaned = text.replace("systemsspecifically", "systems specifically")
            cleaned = clean_text(cleaned)
            if cleaned != stripped:
                para.text = cleaned

        for run in para.runs:
            run.font.name = "Times New Roman"
            if para.style.name == "Heading 1":
                run.font.size = Pt(16)
                run.bold = True
            elif para.style.name == "Heading 2":
                run.font.size = Pt(14)
                run.bold = True
            else:
                run.font.size = Pt(12)


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(2)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in para.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
                        if row_index == 0:
                            run.bold = True


def remove_existing_images(doc: Document) -> None:
    for shape in list(doc.inline_shapes):
        parent = shape._inline.getparent()
        grandparent = parent.getparent()
        grandparent.remove(parent)


def find_paragraph(doc: Document, pattern: str):
    regex = re.compile(pattern, re.I)
    for para in doc.paragraphs:
        if regex.search(para.text.strip()):
            return para
    return None


def insert_picture_before(paragraph, image_path: Path, caption: str, width_inches: float) -> None:
    if not image_path.exists():
        return
    pic_para = paragraph.insert_paragraph_before()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.space_before = Pt(6)
    pic_para.paragraph_format.space_after = Pt(6)
    pic_para.add_run().add_picture(str(image_path), width=Inches(width_inches))

    cap_para = paragraph.insert_paragraph_before()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_before = Pt(0)
    cap_para.paragraph_format.space_after = Pt(10)
    run = cap_para.add_run(caption)
    set_run_font(run, size=11)


def delete_existing_figure_captions(doc: Document) -> None:
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if re.match(r"Figure\s+\d+(\.\d+)?:", text, re.I):
            remove_paragraph(para)


def add_chapter_dividers(doc: Document) -> None:
    chapter_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1" and re.match(r"Chapter\s+\d+:", p.text.strip(), re.I)]
    for para in chapter_paras:
        match = re.match(r"Chapter\s+(\d+):\s*(.+)", para.text.strip(), re.I)
        if not match:
            continue
        number, title = match.groups()
        divider = para.insert_paragraph_before()
        divider.style = doc.styles["Normal"]
        divider.paragraph_format.page_break_before = True
        divider.paragraph_format.space_before = Pt(0)
        divider.paragraph_format.space_after = Pt(0)
        divider.paragraph_format.line_spacing = 1.0
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = divider.add_run(f"CHAPTER {number} :")
        set_run_font(run1, bold=True, size=20)

        title_para = para.insert_paragraph_before()
        title_para.style = doc.styles["Normal"]
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(0)
        title_para.paragraph_format.line_spacing = 1.0
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = title_para.add_run(title.upper())
        set_run_font(run2, bold=True, size=30)

        line_para = para.insert_paragraph_before()
        line_para.style = doc.styles["Normal"]
        line_para.paragraph_format.space_before = Pt(0)
        line_para.paragraph_format.space_after = Pt(0)
        line_para.paragraph_format.left_indent = Inches(1.2)
        line_para.paragraph_format.right_indent = Inches(1.2)
        line_para.paragraph_format.line_spacing = 1.0
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_border(line_para, "bottom", size=12, color="000000", space=1)
        line_para.add_run(" ")

        para.paragraph_format.page_break_before = True


def replace_cover_and_toc(doc: Document) -> None:
    paras = doc.paragraphs
    if len(paras) < 5:
        return

    title_para = paras[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(18)
    title_para.paragraph_format.space_after = Pt(18)
    title_para.text = "WORSHIP SONG LIBRARY"
    for run in title_para.runs:
        set_run_font(run, bold=True, size=20)

    inserted = [
        ("UNIVERSITY NAME", 16, True),
        ("UNIVERSITY LOGO", 12, False),
        ("FINAL YEAR PROJECT REPORT", 16, True),
    ]
    current = title_para
    for text, size, bold in reversed(inserted):
        p = current.insert_paragraph_before(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            set_run_font(run, bold=bold, size=size)
        current = p

    info_para = paras[2]
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.clear()
    lines = [
        "Submitted By: Samuel Nazareth",
        "Roll Number: Not Provided",
        "Guide Name: Not Provided",
        "College: Not Provided",
        "Academic Year: 2025-2026",
    ]
    for idx, line in enumerate(lines):
        run = info_para.add_run(line)
        set_run_font(run, size=12)
        if idx < len(lines) - 1:
            run.add_break()

    toc_heading = find_paragraph(doc, r"^Table of Contents$")
    if toc_heading:
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_heading.paragraph_format.page_break_before = True


def add_toc(doc: Document) -> None:
    toc_heading = find_paragraph(doc, r"^Table of Contents$")
    if not toc_heading:
        return
    p = insert_paragraph_after(toc_heading, "")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and update field to generate table of contents."
    fld_sep.append(fld_text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    set_run_font(run, size=11)


def insert_paragraph_after(paragraph, text: str = "", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def add_images(doc: Document) -> None:
    anchor = find_paragraph(doc, r"5\.2")
    if anchor:
        insert_picture_before(anchor, DIAGRAMS / MANDATORY_IMAGES["System Architecture"][0], MANDATORY_IMAGES["System Architecture"][1], 6.0)

    er_anchor = find_paragraph(doc, r"6\.3 Table Descriptions")
    if er_anchor:
        insert_picture_before(er_anchor, DIAGRAMS / MANDATORY_IMAGES["ER Diagram"][0], MANDATORY_IMAGES["ER Diagram"][1], 6.3)

    chord_anchor = find_paragraph(doc, r"7\.3 Chord Mapping Diagram")
    if chord_anchor:
        insert_picture_before(chord_anchor, DIAGRAMS / MANDATORY_IMAGES["Chord Mapping"][0], MANDATORY_IMAGES["Chord Mapping"][1], 5.7)
        insert_picture_before(chord_anchor, DIAGRAMS / OPTIONAL_IMAGES["Chord Screenshot"][0], OPTIONAL_IMAGES["Chord Screenshot"][1], 2.2)

    render_anchor = find_paragraph(doc, r"8\.2\.1 Rendering Split Diagram")
    if render_anchor:
        insert_picture_before(render_anchor, DIAGRAMS / MANDATORY_IMAGES["Rendering Flow"][0], MANDATORY_IMAGES["Rendering Flow"][1], 5.9)
        insert_picture_before(render_anchor, DIAGRAMS / OPTIONAL_IMAGES["Alignment Comparison"][0], OPTIONAL_IMAGES["Alignment Comparison"][1], 5.8)

    results_anchor = find_paragraph(doc, r"12\.1 User-Facing Outcomes|12\.2 Rendering Speed|Chapter 12:")
    if results_anchor:
        insert_picture_before(results_anchor, DIAGRAMS / OPTIONAL_IMAGES["Responsive Mobile"][0], OPTIONAL_IMAGES["Responsive Mobile"][1], 2.0)


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def w_tag(tag: str) -> str:
    return f"{{{W_NS}}}{tag}"


def build_pg_borders() -> ET.Element:
    pg = ET.Element(w_tag("pgBorders"))
    pg.set(w_tag("offsetFrom"), "page")
    pg.set(w_tag("display"), "allPages")
    pg.set(w_tag("zOrder"), "front")
    for edge in ("top", "left", "bottom", "right"):
        el = ET.SubElement(pg, w_tag(edge))
        el.set(w_tag("val"), "single")
        el.set(w_tag("sz"), "12")
        el.set(w_tag("space"), "20")
        el.set(w_tag("color"), "000000")
    return pg


def ensure_section_format(sect_pr: ET.Element, v_align: str) -> None:
    v_el = sect_pr.find(w_tag("vAlign"))
    if v_el is None:
        v_el = ET.SubElement(sect_pr, w_tag("vAlign"))
    v_el.set(w_tag("val"), v_align)

    pg = sect_pr.find(w_tag("pgBorders"))
    if pg is not None:
        sect_pr.remove(pg)
    sect_pr.append(build_pg_borders())


def clone_section(template: ET.Element, v_align: str, next_page: bool) -> ET.Element:
    sect = deepcopy(template)
    type_el = sect.find(w_tag("type"))
    if next_page:
        if type_el is None:
            type_el = ET.SubElement(sect, w_tag("type"))
        type_el.set(w_tag("val"), "nextPage")
    elif type_el is not None:
        sect.remove(type_el)
    ensure_section_format(sect, v_align)
    return sect


def get_paragraph_text_xml(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(".//w:t", NS)).strip()


def get_paragraph_style_xml(paragraph: ET.Element) -> str:
    ppr = paragraph.find("w:pPr", NS)
    if ppr is None:
        return ""
    pstyle = ppr.find("w:pStyle", NS)
    if pstyle is None:
        return ""
    return pstyle.get(w_tag("val"), "")


def ensure_ppr(paragraph: ET.Element) -> ET.Element:
    ppr = paragraph.find(w_tag("pPr"))
    if ppr is None:
        ppr = ET.Element(w_tag("pPr"))
        paragraph.insert(0, ppr)
    return ppr


def add_section_break_to_paragraph(paragraph: ET.Element, template: ET.Element, v_align: str) -> None:
    ppr = ensure_ppr(paragraph)
    existing = ppr.find(w_tag("sectPr"))
    if existing is not None:
        ppr.remove(existing)
    ppr.append(clone_section(template, v_align, next_page=True))


def patch_docx_layout(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as zin:
        files = {name: zin.read(name) for name in zin.namelist()}

    root = ET.fromstring(files["word/document.xml"])
    body = root.find("w:body", NS)
    if body is None:
        return

    final_sect = body.find("w:sectPr", NS)
    if final_sect is None:
        return

    ensure_section_format(final_sect, "top")
    template = deepcopy(final_sect)

    paragraphs = body.findall("w:p", NS)
    chapter_label_indices = []
    for idx, p in enumerate(paragraphs):
        text = get_paragraph_text_xml(p)
        if re.fullmatch(r"CHAPTER\s+\d+\s*:", text):
            chapter_label_indices.append(idx)

    for label_idx in chapter_label_indices:
        if label_idx <= 0:
            continue
        prev_idx = label_idx - 1
        heading_idx = None
        for scan in range(label_idx + 1, len(paragraphs)):
            text = get_paragraph_text_xml(paragraphs[scan])
            style = get_paragraph_style_xml(paragraphs[scan])
            if style == "Heading1" and re.match(r"Chapter\s+\d+:", text, re.I):
                heading_idx = scan
                break
        if heading_idx is None or heading_idx - 1 <= label_idx:
            continue

        add_section_break_to_paragraph(paragraphs[prev_idx], template, "top")
        add_section_break_to_paragraph(paragraphs[heading_idx - 1], template, "center")

    body_paras = body.findall("w:p", NS)
    for idx, p in enumerate(body_paras):
        text = get_paragraph_text_xml(p)
        if re.fullmatch(r"CHAPTER\s+\d+\s*:", text):
            ppr = ensure_ppr(p)
            jc = ppr.find(w_tag("jc"))
            if jc is None:
                jc = ET.SubElement(ppr, w_tag("jc"))
            jc.set(w_tag("val"), "center")
            spacing = ppr.find(w_tag("spacing"))
            if spacing is None:
                spacing = ET.SubElement(ppr, w_tag("spacing"))
            spacing.set(w_tag("before"), "0")
            spacing.set(w_tag("after"), "0")
            spacing.set(w_tag("line"), "240")
            spacing.set(w_tag("lineRule"), "auto")

            title_p = body_paras[idx + 1] if idx + 1 < len(body_paras) else None
            if title_p is not None:
                tppr = ensure_ppr(title_p)
                tjc = tppr.find(w_tag("jc"))
                if tjc is None:
                    tjc = ET.SubElement(tppr, w_tag("jc"))
                tjc.set(w_tag("val"), "center")
                tspacing = tppr.find(w_tag("spacing"))
                if tspacing is None:
                    tspacing = ET.SubElement(tppr, w_tag("spacing"))
                tspacing.set(w_tag("before"), "0")
                tspacing.set(w_tag("after"), "0")
                tspacing.set(w_tag("line"), "240")
                tspacing.set(w_tag("lineRule"), "auto")

            line_p = body_paras[idx + 2] if idx + 2 < len(body_paras) else None
            if line_p is not None:
                lppr = ensure_ppr(line_p)
                ljc = lppr.find(w_tag("jc"))
                if ljc is None:
                    ljc = ET.SubElement(lppr, w_tag("jc"))
                ljc.set(w_tag("val"), "center")

    files["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)


def save_as_pdf_via_word(docx_path: Path, pdf_path: Path) -> str:
    import win32com.client  # type: ignore

    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    document = None
    try:
        document = word.Documents.Open(str(docx_path))
        document.TablesOfContents(1).Update() if document.TablesOfContents.Count >= 1 else None
        document.Repaginate()
        pages = document.ComputeStatistics(2)
        document.Save()
        document.SaveAs(str(pdf_path), FileFormat=17)
        return str(pages)
    finally:
        if document is not None:
            document.Close(SaveChanges=True)
        word.Quit()


def main() -> None:
    doc = Document(SOURCE)
    style_document(doc)
    remove_existing_images(doc)
    normalize_paragraphs(doc)
    format_tables(doc)
    delete_existing_figure_captions(doc)
    replace_cover_and_toc(doc)
    add_toc(doc)
    add_chapter_dividers(doc)
    add_images(doc)
    doc.save(OUTPUT)
    patch_docx_layout(OUTPUT)

    pdf_path = OUTPUT.with_suffix(".pdf")
    page_info = "unknown"
    try:
        page_info = save_as_pdf_via_word(OUTPUT, pdf_path)
    except Exception as exc:  # pragma: no cover
        print(f"PDF export skipped: {exc}")
    print(f"Saved DOCX: {OUTPUT}")
    print(f"Saved PDF: {pdf_path if pdf_path.exists() else 'not created'}")
    print(f"Page count: {page_info}")


if __name__ == "__main__":
    main()
