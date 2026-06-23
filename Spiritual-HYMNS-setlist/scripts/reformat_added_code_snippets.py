from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOC_PATH = Path(r"D:\worship-song-library\docs\Choord Transpose Library_with_gantt_code_snippets.docx")
OUT_PATH = Path(r"D:\worship-song-library\docs\Choord Transpose Library_with_gantt_code_snippets_reformatted.docx")


def set_run_font(run) -> None:
    run.font.name = "Times New Roman"
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def main() -> None:
    doc = Document(DOC_PATH)
    touched = 0
    for i, paragraph in enumerate(doc.paragraphs):
        text = " ".join(paragraph.text.strip().split())
        if not text.startswith("Code Snippet "):
            continue
        if i + 1 >= len(doc.paragraphs):
            continue

        code_paragraph = doc.paragraphs[i + 1]
        fmt = code_paragraph.paragraph_format
        fmt.line_spacing = 1.5
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(6)
        fmt.first_line_indent = None

        for run in code_paragraph.runs:
            set_run_font(run)
        touched += 1

    doc.save(OUT_PATH)
    print(f"updated={touched}")
    print(OUT_PATH)


if __name__ == "__main__":
    main()
