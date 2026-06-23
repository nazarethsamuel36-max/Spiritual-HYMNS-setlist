from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOC_PATH = Path(r"D:\worship-song-library\docs\Choord Transpose Library.docx")
IMG_PATH = Path(r"D:\worship-song-library\docs\gannt chart.png")
OUT_PATH = Path(r"D:\worship-song-library\docs\Choord Transpose Library_with_gantt.docx")


def add_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def set_run_font(run, size=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if italic is not None:
        run.italic = italic


def main() -> None:
    doc = Document(DOC_PATH)

    anchor = None
    for i, paragraph in enumerate(doc.paragraphs):
        if " ".join(paragraph.text.strip().split()) == "In evaluating the significance of the study, the project team considered performance, correctness, editorial workflows, multilingual search behavior, and long-term data hygiene. Those dimensions are interconnected: a search system is only useful if data is structured correctly, and a rendering engine is only trustworthy if ingestion and storage rules maintain reliable positional indices.":
            anchor = paragraph
            break

    if anchor is None:
        raise RuntimeError("Could not locate end of section 1.5")

    image_p = add_paragraph_after(anchor)
    image_p.style = doc.styles["Normal"]
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.space_before = Pt(6)
    image_p.paragraph_format.space_after = Pt(3)
    image_p.paragraph_format.line_spacing = 1.15
    run = image_p.add_run()
    run.add_picture(str(IMG_PATH), width=Inches(5.6))

    caption_p = add_paragraph_after(image_p)
    caption_p.style = doc.styles["Normal"]
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(0)
    caption_p.paragraph_format.space_after = Pt(6)
    caption_p.paragraph_format.line_spacing = 1.15
    caption_run = caption_p.add_run("Figure 1.1 Gantt Chart for Project Planning")
    set_run_font(caption_run, size=11, italic=True)

    explain_p = add_paragraph_after(caption_p)
    explain_p.style = doc.styles["Normal"]
    explain_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    explain_p.paragraph_format.space_before = Pt(0)
    explain_p.paragraph_format.space_after = Pt(6)
    explain_p.paragraph_format.line_spacing = 1.5
    explain_p.paragraph_format.first_line_indent = Inches(0.3)
    explanation = (
        "The Gantt chart summarizes the execution schedule followed for the Worship Song Library project, beginning with problem understanding and proceeding through analysis, design, implementation, integration, testing, and final documentation. Each activity band represents a bounded phase in which a specific engineering concern was addressed before the next dependent stage began. The sequence reflects the actual development flow of the system, where database design and chord-engine logic were completed before responsive interface integration and validation tasks. As a planning artifact, the chart documents how implementation work was staged to reduce overlap between data modeling, rendering logic, and final verification."
    )
    explain_run = explain_p.add_run(explanation)
    set_run_font(explain_run)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
