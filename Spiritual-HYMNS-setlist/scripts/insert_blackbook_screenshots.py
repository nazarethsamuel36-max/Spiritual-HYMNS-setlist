from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOCX_PATH = Path(r"D:\worship-song-library\docs\worship_song_library_blackbook_final_images.docx")
OUT_PATH = Path(r"D:\worship-song-library\docs\worship_song_library_blackbook_final_images_with_explanations.docx")
IMG_DIR = Path(r"D:\worship-song-library\blackbook-diagrams")


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, align, before, after, line_spacing, first_line_inches=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if first_line_inches is None:
        fmt.first_line_indent = None
    else:
        fmt.first_line_indent = Inches(first_line_inches)
    paragraph.alignment = align


def add_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def add_image_and_paragraph(after_paragraph, image_name: str, explanation: str, width: float = 5.1):
    image_path = IMG_DIR / image_name
    if not image_path.exists():
        return after_paragraph

    image_p = add_paragraph_after(after_paragraph)
    format_paragraph(image_p, WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3, line_spacing=1.15, first_line_inches=None)
    run = image_p.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    exp_p = add_paragraph_after(image_p)
    format_paragraph(exp_p, WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, line_spacing=1.5, first_line_inches=0.3)
    exp_run = exp_p.add_run(explanation)
    set_run_font(exp_run, 12)
    return exp_p


def find_paragraph(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Paragraph not found: {text}")


def main():
    doc = Document(DOCX_PATH)

    chapter5_anchor = find_paragraph(doc, "Figure 5.2 Chord-to-Text Mapping Using Character Index Positions")
    chapter6_anchor = find_paragraph(doc, "[IMAGE TO BE INSERTED - VERIFIED DIAGRAM REQUIRED]")

    p = chapter5_anchor
    p = add_image_and_paragraph(
        p,
        "song_rendered_output.png",
        "This screenshot shows the final rendered song view after structured lyric lines and chord occurrences have been fetched from the relational model and transformed into display-ready spans. Chords are not stored as padded text; they are attached to character positions and rendered above the corresponding lyric fragments during HTML generation. The visible alignment confirms that the rendering layer is reconstructing positional metadata at runtime rather than replaying a static preformatted chord sheet. This behavior is produced by the same mapping pipeline that converts database line records into anchored visual blocks before the browser applies responsive layout rules.",
    )
    p = add_image_and_paragraph(
        p,
        "wrapping_logic.png",
        "This image captures the wrapping stage where the client-side layout engine splits an overlong rendered line into smaller visual segments once the available width is exceeded. During this operation, the JavaScript wrapping routine recalculates the local offsets of every chord token so that each chord remains bound to the correct substring after the break point changes. The screenshot therefore reflects dynamic restructuring of one logical line into multiple rendered lines without mutating the stored lyric text. It directly corresponds to the implementation logic that preserves mapping correctness while adapting to viewport constraints in real time.",
    )
    p = add_image_and_paragraph(
        p,
        "dynamic_transformation.png",
        "This screenshot demonstrates dynamic chord transformation after the rendering structure has already been built, indicating that the system updates musical values while preserving the existing lyric-to-position relationships. The visible change is driven by transposition logic that rewrites chord symbols mathematically from the underlying chord sequence rather than by manually editing displayed text. Because the lyric nodes and chord anchors remain stable, the transformed output retains the same structural alignment after recalculation. The behavior validates the separation between immutable lyric content, stored positional metadata, and runtime chord conversion logic.",
    )

    p = chapter6_anchor
    p = add_image_and_paragraph(
        p,
        "responsive_consistency.png",
        "This screenshot shows the same rendered song content adapting across different viewport widths while maintaining the relative relationship between lyric segments and their associated chord markers. Internally, the browser receives the same structured data model, but the wrapping algorithm recomputes visual breaks so that alignment survives changes in available horizontal space. The stable chord placement confirms that layout responsiveness operates on anchored positions instead of whitespace formatting assumptions. It therefore serves as evidence that the rendering pipeline remains deterministic across responsive states generated from one backend representation.",
    )
    p = add_image_and_paragraph(
        p,
        "mobile_responsive.png",
        "This mobile view confirms that the structured rendering strategy remains correct even when the display width is constrained to a narrow handheld form factor. The visible chord placement indicates that line segmentation and chord re-anchoring are recalculated for the small screen while the original song structure, line order, and positional indices remain unchanged in storage. Rather than compressing a desktop chord sheet, the system rebuilds the visual arrangement from normalized line and chord data for the target device width. The screenshot therefore validates cross-device consistency between database mapping, servlet output, and client-side responsive rendering logic.",
    )
    p = add_image_and_paragraph(
        p,
        "complex_alignment.png",
        "This screenshot illustrates a more difficult alignment case in which dense chord placement and irregular lyric segmentation must still resolve into a readable synchronized output. The system handles such edge cases by processing each chord occurrence as indexed metadata and then redistributing those anchors when the line is wrapped or reflowed for the current viewport. What is visible is not a manually tuned exception, but the result of generalized mapping logic applied to a higher-complexity input pattern. The preserved correspondence between chords and lyric fragments demonstrates that the implementation can absorb alignment stress without falling back to fixed-width text formatting.",
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
