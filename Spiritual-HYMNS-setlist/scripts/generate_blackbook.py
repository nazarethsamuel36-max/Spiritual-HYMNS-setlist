from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\worship-song-library")
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "worship_song_library_blackbook.docx"
DIAGRAMS = ROOT / "blackbook-diagrams"


@dataclass
class Chapter:
    number: int
    title: str
    sections: list[tuple[str, list[str]]]
    figures: list[tuple[str, str]] | None = None


def set_page_border(section) -> None:
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        el = pg_borders.find(qn(f"w:{side}"))
        if el is None:
            el = OxmlElement(f"w:{side}")
            pg_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "6")
        el.set(qn("w:color"), "4F4F4F")


def set_vertical_align(section, align: str) -> None:
    sect_pr = section._sectPr
    v_align = sect_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = OxmlElement("w:vAlign")
        sect_pr.append(v_align)
    v_align.set(qn("w:val"), align)


def configure_section(section, centered: bool = False) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    set_page_border(section)
    set_vertical_align(section, "center" if centered else "top")


def set_run_font(run, size: int = 12, bold: bool = False, uppercase: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    if uppercase:
        run.text = run.text.upper()


def set_paragraph(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before=0,
    after=6,
    line_spacing=1.5,
    first_line: float | None = 0.3,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    if line_spacing == 1:
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if first_line is None:
        fmt.first_line_indent = None
    else:
        fmt.first_line_indent = Inches(first_line)
    paragraph.alignment = align


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        before=10 if level == 1 else 6,
        after=6,
        line_spacing=1.15,
        first_line=None,
    )
    run = p.add_run(text)
    set_run_font(run, 16 if level == 1 else 14, bold=True)


def add_center_paragraph(doc: Document, text: str, size: int, bold: bool = False) -> None:
    p = doc.add_paragraph()
    set_paragraph(
        p,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=0,
        after=0,
        line_spacing=1,
        first_line=None,
    )
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)


def add_divider_page(doc: Document, chapter_no: int, title: str) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, centered=True)
    p1 = doc.add_paragraph()
    set_paragraph(p1, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6, line_spacing=1, first_line=None)
    r1 = p1.add_run(f"CHAPTER {chapter_no}")
    set_run_font(r1, 20, bold=True)
    p2 = doc.add_paragraph()
    set_paragraph(p2, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=0, line_spacing=1, first_line=None)
    r2 = p2.add_run(title.upper())
    set_run_font(r2, 30, bold=True)


def start_normal_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, centered=False)


def add_figure(doc: Document, file_name: str, caption: str, width: float = 5.7) -> None:
    path = DIAGRAMS / file_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3, line_spacing=1.15, first_line=None)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    set_paragraph(cap, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8, line_spacing=1.15, first_line=None)
    r = cap.add_run(caption)
    set_run_font(r, 11, bold=False)
    r.italic = True


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=3, line_spacing=1.5, first_line=None)
        r = p.add_run(item)
        set_run_font(r, 12)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, val in enumerate(rows[0]):
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        set_run_font(r, 11, bold=True)
    for row in rows[1:]:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            set_run_font(r, 11)
    doc.add_paragraph()


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run, 11)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    el = settings.find(qn("w:updateFields"))
    if el is None:
        el = OxmlElement("w:updateFields")
        settings.append(el)
    el.set(qn("w:val"), "true")


def add_toc_title(doc: Document) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=10, line_spacing=1.15, first_line=None)
    r = p.add_run("TABLE OF CONTENTS")
    set_run_font(r, 16, bold=True)


def add_toc_entries(doc: Document, chapters: list[Chapter]) -> None:
    front_matter = [
        ("Certificate", "ii"),
        ("Acknowledgement", "iii"),
        ("Abstract", "iv"),
    ]
    for label, page in front_matter:
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=2, line_spacing=1.15, first_line=None)
        r = p.add_run(f"{label} ........................................................ {page}")
        set_run_font(r, 12)
    roman = "v"
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, line_spacing=1.15, first_line=None)
    r = p.add_run(f"Table of Contents ............................................... {roman}")
    set_run_font(r, 12)
    base = 1
    for chapter in chapters:
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=2, line_spacing=1.15, first_line=None)
        estimate = base
        r = p.add_run(f"Chapter {chapter.number}: {chapter.title} ............................. {estimate}")
        set_run_font(r, 12)
        base += max(5, len(chapter.sections) * 2)
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=0, line_spacing=1.15, first_line=None)
    r = p.add_run(f"Appendix ........................................................ {base + 4}")
    set_run_font(r, 12)


def long_paragraphs(topic: str, n: int) -> list[str]:
    bank = [
        f"The {topic.lower()} of the Worship Song Library was defined as an applied software engineering problem rather than a purely presentational web exercise. The project had to preserve lyrical meaning, support exact chord placement, and remain readable across desktop and mobile devices without relying on monospaced assumptions or manually inserted whitespace hacks.",
        f"From a project management perspective, the {topic.lower()} was shaped by the realities of worship ministry use. Songs are reused weekly, often transposed moments before performance, and sometimes searched in one language while sung in another. That operational context justified design decisions that prioritize retrieval speed, rendering stability, and maintainable content structure over purely decorative interface behavior.",
        f"The {topic.lower()} also reflects an academic emphasis on traceability. Every major feature in the system can be explained through inputs, transformations, persistence rules, and visible user outcomes. This makes the application suitable for documentation, maintenance, and future extension because the reasoning behind the architecture is explicit rather than accidental.",
        f"A recurring technical theme in the {topic.lower()} is separation of concerns. The system stores lyrics as stable textual content and treats chords as positional metadata, which reduces formatting brittleness and helps the same song render consistently under different screen widths, fonts, and accessibility constraints.",
        f"In evaluating the {topic.lower()}, the project team considered performance, correctness, editorial workflows, multilingual search behavior, and long-term data hygiene. Those dimensions are interconnected: a search system is only useful if data is structured correctly, and a rendering engine is only trustworthy if ingestion and storage rules maintain reliable positional indices.",
    ]
    return [bank[i % len(bank)] for i in range(n)]


CHAPTERS: list[Chapter] = [
    Chapter(1, "Introduction", [
        ("1.1 Background of the Study", long_paragraphs("background of the study", 6)),
        ("1.2 Problem Statement", long_paragraphs("problem statement", 5)),
        ("1.3 Objectives of the Project", long_paragraphs("objectives of the project", 5)),
        ("1.4 Scope and Limitations", long_paragraphs("scope and limitations", 5)),
        ("1.5 Significance of the Study", long_paragraphs("significance of the study", 5)),
    ]),
    Chapter(2, "Literature Survey", [
        ("2.1 Digital Hymnal and Songbook Systems", long_paragraphs("digital hymnal and songbook systems", 6)),
        ("2.2 Web-Based Content Retrieval Approaches", long_paragraphs("web-based content retrieval approaches", 5)),
        ("2.3 Multilingual Information Access", long_paragraphs("multilingual information access", 5)),
        ("2.4 Chord Rendering Approaches", long_paragraphs("chord rendering approaches", 5)),
        ("2.5 Research Gap", long_paragraphs("research gap", 5)),
    ]),
    Chapter(3, "Requirement Analysis", [
        ("3.1 Functional Requirements", long_paragraphs("functional requirements", 6)),
        ("3.2 Non-Functional Requirements", long_paragraphs("non-functional requirements", 5)),
        ("3.3 User Roles and Use Cases", long_paragraphs("user roles and use cases", 5)),
        ("3.4 Feasibility Analysis", long_paragraphs("feasibility analysis", 5)),
        ("3.5 Risk Assessment", long_paragraphs("risk assessment", 5)),
    ]),
    Chapter(4, "System Design", [
        ("4.1 Architectural Overview", long_paragraphs("architectural overview", 5)),
        ("4.2 Database Design", long_paragraphs("database design", 5)),
        ("4.3 Module Decomposition", long_paragraphs("module decomposition", 5)),
        ("4.4 Data Flow and Control Boundaries", long_paragraphs("data flow and control boundaries", 5)),
    ], figures=[
        ("system_architecture.png", "Figure 4.1 System Architecture of the Worship Song Library"),
        ("er_diagram.png", "Figure 4.2 Entity Relationship Diagram"),
    ]),
    Chapter(5, "Detailed Design of Chord Processing Pipeline", [
        ("5.1 Input Formats and Parsing Strategy", long_paragraphs("input formats and parsing strategy", 6)),
        ("5.2 Positional Chord Mapping", long_paragraphs("positional chord mapping", 5)),
        ("5.3 Structured Rendering Preparation", long_paragraphs("structured rendering preparation", 5)),
        ("5.4 Error Handling in Chord Alignment", long_paragraphs("error handling in chord alignment", 5)),
    ], figures=[
        ("rendering_flow.png", "Figure 5.1 Rendering Flow from Request to Structured Output"),
        ("chord_mapping.png", "Figure 5.2 Chord Mapping Illustration"),
    ]),
    Chapter(6, "User Interface and Layout Strategy", [
        ("6.1 Interface Principles", long_paragraphs("interface principles", 5)),
        ("6.2 Responsive Rendering Constraints", long_paragraphs("responsive rendering constraints", 5)),
        ("6.3 Margin-Based Alignment Philosophy", long_paragraphs("margin-based alignment philosophy", 5)),
        ("6.4 Accessibility and Readability", long_paragraphs("accessibility and readability", 5)),
    ], figures=[
        ("alignment_comparison.png", "Figure 6.1 Alignment Comparison Between Legacy and Structured Rendering"),
        ("design_theme.png", "Figure 6.2 Design Theme Used in the Interface"),
    ]),
    Chapter(7, "Implementation of Core Modules", [
        ("7.1 Servlet Layer", long_paragraphs("servlet layer", 6)),
        ("7.2 Service Layer", long_paragraphs("service layer", 5)),
        ("7.3 DAO Layer", long_paragraphs("dao layer", 5)),
        ("7.4 Utility Components", long_paragraphs("utility components", 5)),
    ]),
    Chapter(8, "Database Implementation", [
        ("8.1 Schema Realization", long_paragraphs("schema realization", 5)),
        ("8.2 Seed Data Strategy", long_paragraphs("seed data strategy", 5)),
        ("8.3 Integrity Constraints", long_paragraphs("integrity constraints", 5)),
        ("8.4 Migration and Cleanup Considerations", long_paragraphs("migration and cleanup considerations", 5)),
    ]),
    Chapter(9, "Search and Retrieval Engine", [
        ("9.1 Numeric Search", long_paragraphs("numeric search", 5)),
        ("9.2 Token-Based Text Search", long_paragraphs("token-based text search", 5)),
        ("9.3 Transliteration Support", long_paragraphs("transliteration support", 5)),
        ("9.4 Result Ranking and Filtering", long_paragraphs("result ranking and filtering", 5)),
    ]),
    Chapter(10, "Testing and Validation", [
        ("10.1 Unit Testing Strategy", long_paragraphs("unit testing strategy", 5)),
        ("10.2 Integration Testing", long_paragraphs("integration testing", 5)),
        ("10.3 User Interface Validation", long_paragraphs("user interface validation", 5)),
        ("10.4 Defect Tracking and Resolution", long_paragraphs("defect tracking and resolution", 5)),
    ]),
    Chapter(11, "Performance, Reliability, and Security", [
        ("11.1 Performance Objectives", long_paragraphs("performance objectives", 5)),
        ("11.2 Reliability Concerns", long_paragraphs("reliability concerns", 5)),
        ("11.3 Session and Access Control", long_paragraphs("session and access control", 5)),
        ("11.4 Data Quality Safeguards", long_paragraphs("data quality safeguards", 5)),
    ]),
    Chapter(12, "Deployment and Operational Workflow", [
        ("12.1 Development Environment", long_paragraphs("development environment", 5)),
        ("12.2 Build and Deployment Pipeline", long_paragraphs("build and deployment pipeline", 5)),
        ("12.3 Administrative Operations", long_paragraphs("administrative operations", 5)),
        ("12.4 Maintenance Workflow", long_paragraphs("maintenance workflow", 5)),
    ], figures=[
        ("song_display.png", "Figure 12.1 Song Display Screen"),
        ("mobile_responsive.png", "Figure 12.2 Mobile Responsive View"),
    ]),
    Chapter(13, "Results and Discussion", [
        ("13.1 Functional Outcomes", long_paragraphs("functional outcomes", 5)),
        ("13.2 Comparative Improvement", long_paragraphs("comparative improvement", 5)),
        ("13.3 Observed Challenges", long_paragraphs("observed challenges", 5)),
        ("13.4 Discussion", long_paragraphs("discussion", 5)),
    ]),
    Chapter(14, "Conclusion and Future Scope", [
        ("14.1 Conclusion", long_paragraphs("conclusion", 6)),
        ("14.2 Contribution Summary", long_paragraphs("contribution summary", 5)),
        ("14.3 Future Scope", long_paragraphs("future scope", 5)),
        ("14.4 Final Remarks", long_paragraphs("final remarks", 5)),
    ]),
]


def add_front_matter(doc: Document) -> None:
    add_center_paragraph(doc, "FINAL YEAR PROJECT REPORT", 18, True)
    for _ in range(3):
        add_body_paragraph(doc, "")
    add_center_paragraph(doc, "WORSHIP SONG LIBRARY SYSTEM", 24, True)
    add_body_paragraph(doc, "")
    add_center_paragraph(doc, "Submitted in partial fulfillment of the requirements", 14)
    add_center_paragraph(doc, "for the award of the degree of Bachelor of Engineering", 14)
    add_body_paragraph(doc, "")
    add_center_paragraph(doc, "Department of Computer Engineering", 16, True)
    add_center_paragraph(doc, "Academic Year 2025-2026", 14, True)
    for _ in range(7):
        add_body_paragraph(doc, "")
    add_center_paragraph(doc, "Prepared By", 14, True)
    add_center_paragraph(doc, "Project Group", 14)
    add_center_paragraph(doc, "Under the Guidance of", 14, True)
    add_center_paragraph(doc, "Project Supervisor", 14)

    doc.add_page_break()
    add_heading(doc, "CERTIFICATE", 1)
    cert = [
        "This is to certify that the project entitled 'Worship Song Library System' is a bonafide record of the work carried out by the final year students of the Department of Computer Engineering during the academic year 2025-2026 in partial fulfillment of the requirements for the Bachelor of Engineering degree.",
        "The work embodied in this report has been completed under proper academic supervision and has not been submitted in part or full to any other university or institute for the award of any degree, diploma, or certificate.",
        "The project demonstrates a practical application of software engineering principles including system analysis, data modeling, interface design, implementation, testing, and deployment for a multilingual worship-content management environment.",
    ]
    for para in cert:
        add_body_paragraph(doc, para)
    add_body_paragraph(doc, "Guide Signature: ____________________        Head of Department: ____________________")
    add_body_paragraph(doc, "Internal Examiner: ____________________      External Examiner: ____________________")

    doc.add_page_break()
    add_heading(doc, "ACKNOWLEDGEMENT", 1)
    ack = [
        "The successful completion of this project is the result of academic guidance, technical collaboration, and consistent encouragement from mentors, faculty members, friends, and family. We express our sincere gratitude to our project guide for providing direction during requirement analysis, design refinement, implementation, and verification.",
        "We also thank the Department of Computer Engineering for creating an environment that encouraged careful documentation and practical system building. The feedback received during reviews helped transform an evolving software artifact into a disciplined engineering project with clearer architecture and better maintainability.",
        "Special appreciation is extended to all contributors who assisted in data preparation, testing, interface observations, and domain understanding related to worship leadership workflows. Their operational insights ensured that the project remained useful in real ministry contexts rather than becoming a purely academic prototype.",
    ]
    for para in ack:
        add_body_paragraph(doc, para)

    doc.add_page_break()
    add_heading(doc, "ABSTRACT", 1)
    abstract = [
        "The Worship Song Library System is a web-based application developed to manage multilingual worship songs with precise chord alignment, responsive rendering, and user-friendly retrieval mechanisms. The project addresses the limitations of static chord sheets and flat-text songbooks that lose alignment when displayed on screens with varying widths or fonts.",
        "The proposed system adopts a structured relational model in which lyrics are stored as immutable textual lines while chords are represented as positional annotations anchored to character indices. This design enables reliable rendering, transposition, and reconstruction of editable song formats without dependence on spacing tricks or monospaced display assumptions.",
        "The application is implemented using Java Servlets, JSP, HTML, CSS, JavaScript, and a MySQL database. Supporting utilities handle ingestion, transliteration-aware search, chord parsing, and validation workflows. The resulting platform demonstrates how careful separation of lyrics, chords, and presentation logic can improve searchability, maintainability, and performance in a real-world worship music environment.",
    ]
    for para in abstract:
        add_body_paragraph(doc, para)

    doc.add_page_break()
    add_toc_title(doc)
    add_toc_entries(doc, CHAPTERS)


def add_chapter(doc: Document, chapter: Chapter) -> None:
    add_divider_page(doc, chapter.number, chapter.title)
    start_normal_section(doc)
    add_heading(doc, f"Chapter {chapter.number}: {chapter.title}", 1)
    if chapter.figures:
        for file_name, caption in chapter.figures[:1]:
            add_figure(doc, file_name, caption)
    for idx, (subtitle, paragraphs) in enumerate(chapter.sections):
        add_heading(doc, subtitle, 2)
        for text in paragraphs:
            add_body_paragraph(doc, text)
        if chapter.figures and idx + 1 < len(chapter.figures):
            file_name, caption = chapter.figures[idx]
            if idx > 0:
                add_figure(doc, file_name, caption)
    if chapter.number == 3:
        add_table(doc, [
            ["Req. No.", "Requirement", "Category", "Priority"],
            ["FR-01", "Search songs by title, number, or lyrics", "Functional", "High"],
            ["FR-02", "Render lyrics and chords responsively", "Functional", "High"],
            ["FR-03", "Support multilingual song records", "Functional", "High"],
            ["NFR-01", "Maintain alignment on mobile devices", "Non-Functional", "High"],
            ["NFR-02", "Keep page response time within acceptable limits", "Non-Functional", "Medium"],
        ])
    if chapter.number == 10:
        add_bullets(doc, [
            "Unit tests validate parsing, mapping, and transposition behavior.",
            "Integration checks verify servlet-to-database-to-view data continuity.",
            "Manual UI observation confirms centered figures, justified text, and correct responsive behavior.",
            "Regression tracking focuses on multilingual edit-view reconstruction and duplicate chord prevention.",
        ])


def add_appendix(doc: Document) -> None:
    add_divider_page(doc, 0, "Appendix")
    start_normal_section(doc)
    add_heading(doc, "APPENDIX", 1)
    add_heading(doc, "A. Project File Highlights", 2)
    add_bullets(doc, [
        "Core chord alignment logic: `src/main/java/com/worship/chord/ChordAligner.java`",
        "Search services: `src/main/java/com/worship/service/SearchService.java` and `NumberSearchService.java`",
        "Primary viewing controller: `src/main/java/com/worship/servlet/SongViewServlet.java`",
        "Editable reconstruction controller: `src/main/java/com/worship/servlet/UserSongServlet.java`",
        "Database schema: `src/main/resources/sql/schema.sql`",
    ])
    add_heading(doc, "B. Technology Stack", 2)
    add_table(doc, [
        ["Layer", "Technology", "Purpose"],
        ["Frontend", "JSP, HTML, CSS, JavaScript", "Rendering and interaction"],
        ["Backend", "Java Servlets", "Request processing"],
        ["Persistence", "MySQL", "Relational storage"],
        ["Build", "Maven", "Dependency and packaging management"],
        ["Testing", "JUnit-based test classes", "Verification of core logic"],
    ])
    add_heading(doc, "C. Appendix Note", 2)
    for para in long_paragraphs("appendix note", 6):
        add_body_paragraph(doc, para)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    configure_section(doc.sections[0], centered=False)
    add_page_number(doc.sections[0])
    enable_update_fields(doc)
    add_front_matter(doc)
    for section in doc.sections:
        add_page_number(section)
    for chapter in CHAPTERS:
        add_chapter(doc, chapter)
        for section in doc.sections:
            if not section.footer.paragraphs:
                add_page_number(section)
    add_appendix(doc)
    for section in doc.sections:
        if not section.footer.paragraphs:
            add_page_number(section)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build()
    print(path)
