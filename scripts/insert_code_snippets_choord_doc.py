from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOC_PATH = Path(r"D:\worship-song-library\docs\Choord Transpose Library_with_gantt.docx")
OUT_PATH = Path(r"D:\worship-song-library\docs\Choord Transpose Library_with_gantt_code_snippets.docx")


SNIPPETS = [
    (
        "5.1 Input Formats and Parsing Strategy",
        1,
        "Code Snippet 5.1: Bracket Chord Parsing",
        """String input = normalizeInput(rawLine);
StringBuilder lyrics = new StringBuilder();
List<ChordOccurrence> chords = new ArrayList<>();
int cursor = 0;
while (i < input.length()) {
    if (input.charAt(i) == '[') {
        int closingIndex = input.indexOf(']', i);
        String chordText = input.substring(i + 1, closingIndex).trim();
        chords.add(new ChordOccurrence(rebuild(parseChordToken(chordText)), cursor));
    }
}""",
    ),
    (
        "5.2 Positional Chord Mapping",
        1,
        "Code Snippet 5.2: Position-Based Chord Mapping",
        """List<ChordOccurrence> result = new ArrayList<>();
for (ChordOccurrence ext : extLine.getChords()) {
    int pos = Math.max(0, Math.min(ext.getPosition(), dbText.length()));
    result.add(new ChordOccurrence(ext.getChord(), pos));
}
Collections.sort(result);
return result;""",
    ),
    (
        "5.3 Structured Rendering Preparation",
        1,
        "Code Snippet 5.3: Building Structured Lines",
        """for (Section section : sections) {
    structuredLines.add(new StructuredLine("{" + section.getLabel() + "}"));
    for (SongLine line : section.getLines()) {
        StructuredLine sl = ChordParser.parseStructuredLine(line.getText());
        List<ChordOccurrence> lineChords = chordMap.get(line.getId());
        if (lineChords != null && !lineChords.isEmpty()) sl.setChords(lineChords);
        structuredLines.add(sl);
    }
}""",
    ),
    (
        "7.1 Servlet Layer",
        1,
        "Code Snippet 7.1: Search Servlet Request Flow",
        """String query = request.getParameter("q");
String pageParam = request.getParameter("page");
if (query == null || query.trim().isEmpty()) {
    request.getRequestDispatcher("/jsp/search.jsp").forward(request, response);
    return;
}
Map<String, Object> serviceResult = searchService.search(query.trim(), pageNum, pageSize);
request.setAttribute("searchResults", serviceResult.get("results"));
request.getRequestDispatcher("/jsp/search.jsp").forward(request, response);""",
    ),
    (
        "7.3 DAO Layer",
        1,
        "Code Snippet 7.2: DAO Query with PreparedStatement",
        """String sql = "SELECT * FROM songs WHERE song_number = ? AND " + getVisibilityFilter();
try (Connection conn = DBConnection.getConnection();
     PreparedStatement ps = conn.prepareStatement(sql)) {
    ps.setInt(1, songNumber);
    try (ResultSet rs = ps.executeQuery()) {
        while (rs.next()) {
            Song song = mapResultSetToSong(rs);
            songs.add(song);
        }
    }
}""",
    ),
    (
        "8.1 Schema Realization",
        1,
        "Code Snippet 8.1: line_chords Table Schema",
        """CREATE TABLE IF NOT EXISTS line_chords (
    id         INT PRIMARY KEY AUTO_INCREMENT,
    line_id    INT NOT NULL,
    chord      VARCHAR(30) NOT NULL,
    char_index INT NOT NULL,
    confidence DECIMAL(4,3),
    flag       VARCHAR(30),
    FOREIGN KEY (line_id) REFERENCES song_lines(id) ON DELETE CASCADE
);""",
    ),
    (
        "9.2 Token-Based Text Search",
        1,
        "Code Snippet 9.1: Token-Based Search Pipeline",
        """String normalizedQuery = normalizeQuery(query);
String[] tokens = tokenizeQuery(normalizedQuery);
Map<String, List<String>> variantMap = buildVariantMap(tokens);
List<Song> candidates = songDAO.searchSongsByTokensAndVariants(Arrays.asList(tokens), variantMap);
List<Song> filtered = candidates.stream()
    .filter(song -> matchesAllTokens(song, tokens, variantMap))
    .collect(Collectors.toList());
SongScore score = calculateSongScore(song, tokens, variantMap);""",
    ),
]


def add_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def set_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def style_label(paragraph, text):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.0
    pf.first_line_indent = None
    run = paragraph.add_run(text)
    set_font(run, size=11, bold=True, italic=True)


def style_code(paragraph, code_text):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    pf.first_line_indent = None
    run = paragraph.add_run(code_text)
    set_font(run, name="Courier New", size=9)


def main():
    doc = Document(DOC_PATH)

    for heading_text, para_offset, label, code in reversed(SNIPPETS):
        idx = None
        for i, p in enumerate(doc.paragraphs):
            if " ".join(p.text.strip().split()) == heading_text:
                idx = i
                break
        if idx is None:
            raise RuntimeError(f"Heading not found: {heading_text}")

        anchor = doc.paragraphs[idx + para_offset]
        label_p = add_paragraph_after(anchor)
        style_label(label_p, label)
        code_p = add_paragraph_after(label_p)
        style_code(code_p, code)

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
