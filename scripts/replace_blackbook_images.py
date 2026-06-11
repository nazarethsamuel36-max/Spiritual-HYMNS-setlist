from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"D:\worship-song-library")
BASE = ROOT / "docs" / "worship_song_library_blackbook.docx"
OUT = ROOT / "docs" / "worship_song_library_blackbook_corrected.docx"
IMG_DIR = ROOT / "docs" / "blackbook" / "extracted_media"

PLACEHOLDER = "[IMAGE TO BE INSERTED - VERIFIED DIAGRAM REQUIRED]"


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6, line_spacing=1.15, first_line=None):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.first_line_indent = first_line
    p.alignment = align


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return paragraph._parent.add_paragraph()._p  # placeholder; replaced below


def add_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, paragraph._parent)


def replace_image_paragraph(paragraph, image_path: Path, width: float):
    clear_paragraph(paragraph)
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3, line_spacing=1.15)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))


def replace_with_placeholder(paragraph, caption_paragraph):
    clear_paragraph(paragraph)
    format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3, line_spacing=1.15)
    run = paragraph.add_run(PLACEHOLDER)
    set_run_font(run, 12, italic=True)
    clear_paragraph(caption_paragraph)
    format_paragraph(caption_paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8, line_spacing=1.15)


def add_figure_after(caption_paragraph, image_path: Path, caption: str, width: float):
    image_p = add_paragraph_after(caption_paragraph)
    format_paragraph(image_p, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=3, line_spacing=1.15)
    run = image_p.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    cap_p = add_paragraph_after(image_p)
    format_paragraph(cap_p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8, line_spacing=1.15)
    cap_run = cap_p.add_run(caption)
    set_run_font(cap_run, 11, italic=True)


def main():
    doc = Document(BASE)
    paragraphs = doc.paragraphs

    def by_text(text: str):
        for idx, p in enumerate(paragraphs):
            if p.text.strip() == text:
                return idx, p
        raise ValueError(text)

    idx_41, p_41 = by_text("Figure 4.1 System Architecture of the Worship Song Library")
    replace_image_paragraph(paragraphs[idx_41 - 1], IMG_DIR / "image6.png", 5.7)
    add_figure_after(
        p_41,
        IMG_DIR / "image7.png",
        "Figure 4.2 Entity Relationship Diagram of the Worship Song Library Database",
        5.5,
    )

    idx_51, p_51 = by_text("Figure 5.1 Rendering Flow from Request to Structured Output")
    replace_image_paragraph(paragraphs[idx_51 - 1], IMG_DIR / "image10.png", 5.7)
    add_figure_after(
        p_51,
        IMG_DIR / "image8.png",
        "Figure 5.2 Chord-to-Text Mapping Using Character Index Positions",
        5.5,
    )

    idx_61, p_61 = by_text("Figure 6.1 Alignment Comparison Between Legacy and Structured Rendering")
    replace_with_placeholder(paragraphs[idx_61 - 1], p_61)

    idx_121, p_121 = by_text("Figure 12.1 Song Display Screen")
    replace_with_placeholder(paragraphs[idx_121 - 1], p_121)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
